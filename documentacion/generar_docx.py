# -*- coding: utf-8 -*-
"""Genera PROYECTO_DANILO_ARAMAYO.docx en formato APA 7ma edicion
(Times New Roman 12, interlineado 1,5) a partir de los capitulos Markdown."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = 'Times New Roman'
FIG_TITULOS = [
    ('Diagrama entidad-relación de la base de datos optica_db', 'img/fig1_er_base_datos.png'),
    ('Flujo de navegación del sistema según el rol del usuario', 'img/fig2_flujo_navegacion.png'),
    ('Arquitectura en capas implementada en la solución', 'img/fig3_arquitectura_capas.png'),
    ('Secuencia de autenticación JWT con control de acceso por roles', 'img/fig4_secuencia_jwt.png'),
    ('Pirámide de pruebas aplicada al proyecto', 'img/fig5_piramide_pruebas.png'),
    ('Arquitectura de despliegue con Docker Compose', 'img/fig6_despliegue_docker.png'),
    ('Canalización conceptual de despliegue del prototipo', 'img/fig7_pipeline.png'),
]
TABLA_TITULOS = [
    'Matriz de consistencia técnica y metodológica del proyecto',
    'Estructura de la tabla roles',
    'Estructura de la tabla usuarios',
    'Estructura de la tabla clientes',
    'Estructura de la tabla recetas',
    'Estructura de la tabla productos',
    'Estructura de la tabla ventas',
    'Estructura de la tabla detalle_venta',
    'Estructura de la tabla orden_trabajo',
    'Estructura de la tabla recibos',
    'Contrato general de la API REST',
    'Mapa de endpoints del sistema',
    'Pantallas del sistema y su contenido principal',
    'Validación de datos en cuatro niveles',
    'Servicios definidos en Docker Compose',
    'Variables de entorno del despliegue',
    'URLs del sistema desplegado',
    'Suites de pruebas unitarias ejecutadas',
    'Casos de prueba funcionales: autenticación y usuarios',
    'Casos de prueba funcionales: clientes y recetas',
    'Casos de prueba funcionales: inventario y ventas',
    'Casos de prueba funcionales: órdenes de trabajo y reportes',
    'Evaluación de los objetivos específicos frente al prototipo',
]

# ---------- utilidades base ----------

def fuente(run, size=12, bold=False, italic=False, mono=False):
    run.font.name = 'Courier New' if mono else TNR
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_field(par, instr, placeholder='...'):
    r1, r2, r3, r4, r5 = (par.add_run() for _ in range(5))
    for r in (r1, r2, r3, r4, r5):
        fuente(r)
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = placeholder
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r1._r.append(b); r2._r.append(i); r3._r.append(s); r4._r.append(t); r5._r.append(e)

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`[^`]+`)')

def runs_md(par, texto, size=12, bold=False, italic=False):
    """Convierte **negrita**, *cursiva* y `codigo` a runs."""
    for trozo in INLINE.split(texto):
        if not trozo:
            continue
        if trozo.startswith('**') and trozo.endswith('**'):
            fuente(par.add_run(trozo[2:-2]), size, bold=True, italic=italic)
        elif trozo.startswith('*') and trozo.endswith('*') and len(trozo) > 2:
            fuente(par.add_run(trozo[1:-1]), size, bold=bold, italic=True)
        elif trozo.startswith('`') and trozo.endswith('`'):
            fuente(par.add_run(trozo[1:-1]), size - 1, bold, False, mono=True)
        else:
            fuente(par.add_run(trozo), size, bold, italic)

def parrafo(doc, align=None, indent_first=False, before=0, after=0, line=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if indent_first:
        pf.first_line_indent = Cm(1.27)
    return p

def heading(doc, texto, nivel):
    estilos = {1: ('Heading 1', WD_ALIGN_PARAGRAPH.CENTER),
               2: ('Heading 2', WD_ALIGN_PARAGRAPH.LEFT),
               3: ('Heading 3', WD_ALIGN_PARAGRAPH.LEFT)}
    nombre, alineacion = estilos[nivel]
    p = parrafo(doc, align=alineacion, before=12 if nivel == 1 else 10,
                after=6 if nivel == 1 else 4)
    p.style = doc.styles[nombre]
    fuente(p.add_run(texto), 12, bold=True, italic=(nivel == 3))
    return p

def cuerpo(doc, texto):
    p = parrafo(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW if hasattr(WD_ALIGN_PARAGRAPH, 'JUSTIFY_LOW') else None,
                indent_first=True, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    runs_md(p, texto)
    return p

def cita_bloque(doc, texto):
    p = parrafo(doc, after=6)
    p.paragraph_format.left_indent = Cm(1.27)
    txt = texto.strip()
    if not (txt.startswith('«') or txt.startswith('"')):
        txt = f'"{txt}"'
    runs_md(p, txt)
    return p

def vineta(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet' if nivel == 0 else 'List Bullet 2')
    pf = p.paragraph_format
    pf.line_spacing = 1.5; pf.space_after = Pt(0)
    runs_md(p, texto)
    return p

def bloque_codigo(doc, lineas):
    for ln in lineas:
        p = parrafo(doc, line=1.0, after=0)
        p.paragraph_format.left_indent = Cm(1.27)
        fuente(p.add_run(ln if ln else ' '), 10, mono=True)
    parrafo(doc, after=6, line=1.0)

# ---------- tablas APA ----------

def _borders(tbl):
    tblPr = tbl._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for borde, val in (('top', 'single'), ('bottom', 'single'), ('left', 'none'),
                       ('right', 'none'), ('insideH', 'none'), ('insideV', 'none')):
        e = OxmlElement(f'w:{borde}')
        e.set(qn('w:val'), val)
        if val == 'single':
            e.set(qn('w:sz'), '8'); e.set(qn('w:color'), '000000')
        tb.append(e)
    tblPr.append(tb)

def _header_bottom(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tb = OxmlElement('w:tcBorders')
        e = OxmlElement('w:bottom')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '8'); e.set(qn('w:color'), '000000')
        tb.append(e); tcPr.append(tb)

def _caption(doc, etiqueta, titulo, num_ref):
    p1 = parrafo(doc, before=8, after=2)
    fuente(p1.add_run(etiqueta + ' '), 12, bold=True)
    add_field(p1, f'SEQ {etiqueta} \\* ARABIC', str(num_ref))
    p2 = parrafo(doc, after=4)
    fuente(p2.add_run(titulo), 12, italic=True)

def tabla_apa(doc, headers, filas, titulo_idx, nota=None):
    n = tabla_apa.contador
    _caption(doc, 'Tabla', TABLA_TITULOS[titulo_idx], n + 1)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _borders(t)
    hdr = t.rows[0]
    trPr = hdr._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
    for j, htxt in enumerate(headers):
        c = hdr.cells[j]
        p = c.paragraphs[0]; p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runs_md(p, htxt.replace('**', ''), size=11, bold=True)
    _header_bottom(hdr)
    for fila in filas:
        row = t.add_row()
        for j, val in enumerate(fila[:len(headers)]):
            c = row.cells[j]
            p = c.paragraphs[0]; p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(2)
            runs_md(p, str(val).replace('<br/>', ' '), size=11)
    if nota:
        p = parrafo(doc, before=4, after=8)
        fuente(p.add_run('Nota. '), 11, italic=True)
        runs_md(p, nota, size=11)
    else:
        parrafo(doc, after=6, line=1.0)
    tabla_apa.contador += 1
tabla_apa.contador = 0

def figura_apa(doc, idx_fig):
    titulo, ruta = FIG_TITULOS[idx_fig - 1]
    p1 = parrafo(doc, before=8, after=2)
    fuente(p1.add_run('Figura '), 12, bold=True)
    add_field(p1, 'SEQ Figura \\* ARABIC', str(idx_fig))
    p2 = parrafo(doc, after=4)
    fuente(p2.add_run(titulo), 12, italic=True)
    pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(4)
    pic_p.add_run().add_picture(ruta, width=Inches(6.1))

def nota_suelta(doc, texto):
    p = parrafo(doc, before=2, after=8)
    fuente(p.add_run('Nota. '), 11, italic=True)
    runs_md(p, texto.lstrip('\\*').strip(), size=11)

# ---------- indices ----------

def indice(doc, titulo, instr, placeholder):
    heading(doc, titulo, 1)
    p = parrafo(doc)
    add_field(p, instr, placeholder)
    aviso = parrafo(doc, before=4, after=0)
    fuente(aviso.add_run('(Al abrir el documento: seleccione este índice, clic derecho → '
                         'Actualizar campos, o presione F9 para generar el contenido.)'), 10, italic=True)

# ---------- parser markdown ----------

def parse_md(path, doc, estado):
    lineas = open(path, encoding='utf-8').read().split('\n')
    i = 0
    while i < len(lineas):
        ln = lineas[i].rstrip()
        if ln.startswith('```'):
            lang = ln[3:].strip().lower()
            bloque = []
            i += 1
            while i < len(lineas) and not lineas[i].startswith('```'):
                bloque.append(lineas[i]); i += 1
            if lang == 'mermaid':
                estado['fig'] += 1
                figura_apa(doc, estado['fig'])
            else:
                bloque_codigo(doc, bloque)
        elif ln.startswith('|'):
            filas_md = []
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                filas_md.append(lineas[i].strip()); i += 1
            celdas = [[c.strip() for c in f.strip('|').split('|')] for f in filas_md
                      if not re.match(r'^\|[\s:\-|]+\|$', f)]
            if celdas:
                tabla_apa(doc, celdas[0], celdas[1:], titulo_idx=estado['tab'])
                estado['tab'] += 1
            continue
        elif ln.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.*)$', ln)
            nivel = len(m.group(1))
            heading(doc, m.group(2).strip(), min(nivel, 3))
        elif ln.startswith('> '):
            cita_bloque(doc, ln[2:].strip())
        elif ln.lstrip().startswith('- '):
            sangria = len(ln) - len(ln.lstrip())
            vineta(doc, ln.lstrip()[2:].strip(), nivel=0 if sangria < 2 else 1)
        elif re.match(r'^\d+\.\s+', ln):
            cuerpo(doc, ln)
        elif ln.strip().startswith('\\*'):
            nota_suelta(doc, ln.strip())
        elif ln.strip():
            if ln.strip() == '.':
                pass  # marcador vacío del original
            else:
                cuerpo(doc, ln.strip())
        i += 1

# ---------- construccion ----------

doc = Document()

# estilo Normal
normal = doc.styles['Normal']
normal.font.name = TNR; normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), TNR)
npf = normal.paragraph_format
npf.line_spacing = 1.5; npf.space_before = Pt(0); npf.space_after = Pt(0)

for hs, it in (('Heading 1', False), ('Heading 2', False), ('Heading 3', True)):
    st = doc.styles[hs]
    st.font.name = TNR; st.font.size = Pt(12)
    st.font.bold = True; st.font.italic = it
    st.font.color.rgb = RGBColor(0, 0, 0)

sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)

# numero de pagina arriba a la derecha
hdr_p = sec.header.paragraphs[0]
hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_field(hdr_p, 'PAGE', '1')

def linea_centro(texto, bold=False, size=12, antes=0):
    p = parrafo(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=antes)
    fuente(p.add_run(texto), size, bold=bold)
    return p

# ===== PORTADA =====
linea_centro('UNIVERSIDAD UNIOR', bold=True, size=16, antes=36)
linea_centro('INGENIERÍA DE SISTEMAS', bold=True, size=14, antes=18)
for _ in range(5):
    parrafo(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
linea_centro('"SISTEMA DE INFORMACIÓN WEB FULL STACK PARA LA GESTIÓN '
             'INTEGRAL DE LA ÓPTICA VICTORIA"', bold=True, size=14)
for _ in range(4):
    parrafo(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
linea_centro('Estudiante: Aramayo Garisto Danilo', size=13)
parrafo(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
linea_centro('Oruro – Bolivia', size=13)
linea_centro('2026', size=13)
doc.add_page_break()

# ===== INDICES =====
indice(doc, 'ÍNDICE GENERAL', 'TOC \\o "1-3" \\h \\z \\u', 'Índice general del documento')
doc.add_page_break()
indice(doc, 'ÍNDICE DE TABLAS', 'TOC \\h \\z \\c "Tabla"', 'Relación de tablas del documento')
doc.add_page_break()
indice(doc, 'ÍNDICE DE FIGURAS', 'TOC \\h \\z \\c "Figura"', 'Relación de figuras del documento')
doc.add_page_break()

# ===== RESUMEN =====
heading(doc, 'Resumen', 1)
resumen = open('00-resumen.md', encoding='utf-8').read().split('\n')
for ln in resumen:
    ln = ln.rstrip()
    if not ln or ln.startswith('#'):
        continue
    if ln.startswith('**Palabras clave:**'):
        p = parrafo(doc, before=8)
        fuente(p.add_run('Palabras clave: '), 12, bold=True, italic=True)
        runs_md(p, ln.replace('**Palabras clave:**', '').strip())
    else:
        p = parrafo(doc, after=6)
        runs_md(p, ln)
doc.add_page_break()

# ===== CAPITULOS I-VI =====
orden = ['01-introduccion.md', '02-marco-teorico.md', '03-metodologia.md',
         '04-modelado-diseno.md', '05-implementacion.md', '06-conclusiones.md']
estado = {'fig': 0, 'tab': 0}
for cap in orden:
    parse_md(cap, doc, estado)
    doc.add_page_break()

# ===== REFERENCIAS =====
heading(doc, 'Referencias', 1)
refs = [
    'Arjonilla, R. (s.f.). ¿Qué es Backend? Rafarjonilla. https://rafarjonilla.com/que-es/backend/',
    'EBAC. (17 de julio de 2023). Qué es React y para qué sirve. EBAC Blog. https://ebac.mx/blog/que-es-react',
    'Flores, J. L. (4 de septiembre de 2019). Qué es NodeJS y para qué sirve. OpenWebinars. https://openwebinars.net/blog/que-es-nodejs/',
    'IONOS. (6 de julio de 2024). ¿Qué es el frontend? IONOS Digital Guide. https://www.ionos.es/digitalguide/paginas-web/creacion-de-paginas-web/que-es-el-frontend/',
    'Maceda, H. C. (2015). Arquitectura de software.',
    'Mixpanel Team. (6 de diciembre de 2024). What is a tech stack? Mixpanel Blog. https://mixpanel.com/blog/what-is-a-technology-stack/',
]
for ref in refs:
    p = parrafo(doc, after=6)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27); pf.first_line_indent = Cm(-1.27)   # sangria francesa
    fuente(p.add_run(ref), 12)

salida = '../PROYECTO_DANILO_ARAMAYO.docx'
doc.save(salida)
print(f'OK -> {salida}')
print(f'Tablas numeradas: {tabla_apa.contador} | Figuras insertadas: {estado["fig"]}')
